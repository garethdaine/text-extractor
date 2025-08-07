"""Integration tests for the complete text extraction workflow."""


import pytest

from text_extractor import extract_text_from_file
from text_extractor.models import ExtractedText


class TestIntegration:
    """Test complete integration workflow."""

    def test_extract_text_from_txt_file(self, tmp_path):
        """Test extracting text from a TXT file."""
        test_file = tmp_path / "test.txt"
        test_content = "Hello, World!\nThis is a test file."
        test_file.write_text(test_content, encoding="utf-8")

        result = extract_text_from_file(str(test_file))

        assert isinstance(result, ExtractedText)
        assert result.file_type == "txt"
        assert result.ocr_used is False
        assert "Hello, World!" in result.text
        assert "This is a test file." in result.text

    def test_extract_text_from_csv_file(self, tmp_path):
        """Test extracting text from a CSV file."""
        test_file = tmp_path / "test.csv"
        test_content = """Name,Age,City
John Doe,30,New York
Jane Smith,25,Los Angeles"""
        test_file.write_text(test_content, encoding="utf-8")

        result = extract_text_from_file(str(test_file))

        assert isinstance(result, ExtractedText)
        assert result.file_type == "csv"
        assert result.ocr_used is False
        assert "Name" in result.text
        assert "John Doe" in result.text
        assert "Jane Smith" in result.text

    def test_extract_text_from_docx_file(self, tmp_path):
        """Test extracting text from a DOCX file."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        doc.add_paragraph("Hello from DOCX!")
        doc.add_paragraph("This is a test document.")

        test_file = tmp_path / "test.docx"
        doc.save(str(test_file))

        result = extract_text_from_file(str(test_file))

        assert isinstance(result, ExtractedText)
        assert result.file_type == "docx"
        assert result.ocr_used is False
        assert "Hello from DOCX!" in result.text
        assert "This is a test document." in result.text

    def test_extract_text_from_pdf_file(self, tmp_path):
        """Test extracting text from a PDF file."""
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            pytest.skip("reportlab not available")

        test_file = tmp_path / "test.pdf"
        c = canvas.Canvas(str(test_file))
        c.drawString(100, 750, "Hello from PDF!")
        c.drawString(100, 730, "This is a test document.")
        c.save()

        result = extract_text_from_file(str(test_file))

        assert isinstance(result, ExtractedText)
        assert result.file_type == "pdf"
        assert result.ocr_used is False
        assert "Hello from PDF!" in result.text
        assert "This is a test document." in result.text

    def test_extract_text_from_image_file(self, tmp_path):
        """Test extracting text from an image file."""
        try:
            from PIL import Image, ImageDraw
        except ImportError:
            pytest.skip("Pillow not available")

        img = Image.new('RGB', (300, 150), color='white')
        draw = ImageDraw.Draw(img)

        try:
            draw.text((50, 50), "Hello from Image!", fill='black')
        except Exception:
            pass

        test_file = tmp_path / "test.png"
        img.save(str(test_file))

        result = extract_text_from_file(str(test_file))

        assert isinstance(result, ExtractedText)
        assert result.file_type == "png"
        assert result.ocr_used is True
        assert len(result.pages) == 1
        assert result.pages[0].ocr is True

    def test_extract_text_from_unsupported_file(self):
        """Test that unsupported file types raise appropriate errors."""
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text_from_file("document.xyz")

    def test_extract_text_from_nonexistent_file(self):
        """Test that nonexistent files raise appropriate errors."""
        with pytest.raises(FileNotFoundError):
            extract_text_from_file("nonexistent.txt")

    def test_extract_text_with_different_file_extensions(self, tmp_path):
        """Test extracting text from files with different extensions."""
        # Test case sensitivity
        test_content = "Test content"

        # Create files with different case extensions
        files = [
            ("test.TXT", test_content),
            ("test.PDF", ""),  # Will be created as PDF
            ("test.CSV", "Name,Value\nTest,Content"),
        ]

        for filename, content in files:
            test_file = tmp_path / filename

            if filename.endswith('.PDF'):
                try:
                    from reportlab.pdfgen import canvas
                    c = canvas.Canvas(str(test_file))
                    c.drawString(100, 750, content)
                    c.save()
                except ImportError:
                    continue
            else:
                test_file.write_text(content, encoding="utf-8")

            result = extract_text_from_file(str(test_file))
            assert isinstance(result, ExtractedText)
            assert result.file_type in ["txt", "pdf", "csv"]

    def test_extract_text_returns_consistent_structure(self, tmp_path):
        """Test that all parsers return consistent structure."""
        # Test with TXT file
        test_file = tmp_path / "test.txt"
        test_content = "Simple test content"
        test_file.write_text(test_content, encoding="utf-8")

        result = extract_text_from_file(str(test_file))

        # Verify consistent structure
        assert hasattr(result, 'text')
        assert hasattr(result, 'file_type')
        assert hasattr(result, 'ocr_used')
        assert hasattr(result, 'pages')
        assert isinstance(result.text, str)
        assert isinstance(result.file_type, str)
        assert isinstance(result.ocr_used, bool)
        assert isinstance(result.pages, list)

        if result.pages:
            page = result.pages[0]
            assert hasattr(page, 'page_number')
            assert hasattr(page, 'text')
            assert hasattr(page, 'ocr')
            assert isinstance(page.page_number, int)
            assert isinstance(page.text, str)
            assert isinstance(page.ocr, bool)
