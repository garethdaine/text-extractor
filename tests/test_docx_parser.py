"""Tests for DOCX file parser."""

import pytest

from text_extractor.models import ExtractedText
from text_extractor.parsers import docx_parser


class TestDocxParser:
    """Test DOCX file parsing functionality."""

    def test_parse_simple_docx_file(self, tmp_path):
        """Test parsing a simple DOCX file."""
        # Create a simple DOCX file using python-docx
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        doc.add_paragraph("Hello, World!")
        doc.add_paragraph("This is a test document.")

        test_file = tmp_path / "test.docx"
        doc.save(str(test_file))

        result = docx_parser.parse(str(test_file))

        assert isinstance(result, ExtractedText)
        assert result.file_type == "docx"
        assert result.ocr_used is False
        assert "Hello, World!" in result.text
        assert "This is a test document." in result.text
        assert len(result.pages) == 1
        assert result.pages[0].page_number == 1
        assert result.pages[0].ocr is False

    def test_parse_docx_with_multiple_paragraphs(self, tmp_path):
        """Test parsing DOCX with multiple paragraphs."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")

        test_file = tmp_path / "multi.docx"
        doc.save(str(test_file))

        result = docx_parser.parse(str(test_file))

        assert result.file_type == "docx"
        assert result.ocr_used is False
        assert "First paragraph" in result.text
        assert "Second paragraph" in result.text
        assert "Third paragraph" in result.text

    def test_parse_docx_with_special_characters(self, tmp_path):
        """Test parsing DOCX with special characters."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        doc.add_paragraph("Special chars: !@#$%^&*()_+-=[]{}|;':\",./<>?")
        doc.add_paragraph("Unicode: 你好世界 🌍 éñç")

        test_file = tmp_path / "special.docx"
        doc.save(str(test_file))

        result = docx_parser.parse(str(test_file))

        assert result.file_type == "docx"
        assert "Special chars:" in result.text
        assert "!@#$%^&*()" in result.text
        assert "Unicode:" in result.text
        assert "你好世界" in result.text
        assert "🌍" in result.text
        assert "éñç" in result.text

    def test_parse_empty_docx(self, tmp_path):
        """Test parsing an empty DOCX file."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        test_file = tmp_path / "empty.docx"
        doc.save(str(test_file))

        result = docx_parser.parse(str(test_file))

        assert result.file_type == "docx"
        assert result.ocr_used is False
        # Empty DOCX might still have some default content
        assert len(result.pages) == 1
        assert result.pages[0].ocr is False

    def test_parse_docx_with_tables(self, tmp_path):
        """Test parsing DOCX with tables."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        doc.add_paragraph("Before table")

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"

        doc.add_paragraph("After table")

        test_file = tmp_path / "table.docx"
        doc.save(str(test_file))

        result = docx_parser.parse(str(test_file))

        assert result.file_type == "docx"
        assert result.ocr_used is False
        assert "Before table" in result.text
        assert "After table" in result.text
        # Table content should be extracted
        assert "Header 1" in result.text or "Data 1" in result.text

    def test_parse_nonexistent_docx_file(self):
        """Test that parsing a nonexistent DOCX file raises an appropriate error."""
        with pytest.raises(FileNotFoundError):
            docx_parser.parse("nonexistent.docx")

    def test_parse_docx_with_formatted_text(self, tmp_path):
        """Test parsing DOCX with formatted text."""
        try:
            from docx import Document

        except ImportError:
            pytest.skip("python-docx not available")

        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Bold text").bold = True
        p.add_run(" and ")
        p.add_run("italic text").italic = True

        test_file = tmp_path / "formatted.docx"
        doc.save(str(test_file))

        result = docx_parser.parse(str(test_file))

        assert result.file_type == "docx"
        assert result.ocr_used is False
        assert "Bold text" in result.text
        assert "italic text" in result.text
