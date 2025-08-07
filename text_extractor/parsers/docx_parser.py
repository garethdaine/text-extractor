"""DOCX parser implementation using :mod:`python-docx`."""

from ..models import ExtractedText, PageText


def parse(file_path: str) -> ExtractedText:
    """Parse a DOCX file and return extracted text.

    Parameters
    ----------
    file_path: str
        Path to the DOCX file.

    Returns
    -------
    ExtractedText
        Structured text extracted from the DOCX document.
    """
    from docx import Document  # Imported lazily
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = Document(file_path)
    except PackageNotFoundError as e:
        raise FileNotFoundError(f"DOCX file not found: {file_path}") from e

    # Extract text from paragraphs
    text_parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    text = "\n".join(text_parts)
    pages = [PageText(page_number=1, text=text, ocr=False)]
    return ExtractedText(text=text, file_type="docx", pages=pages)
