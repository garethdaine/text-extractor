"""Async DOCX file parser implementation using :mod:`python-docx`."""

import asyncio

from ..models import ExtractedText, PageText


async def parse(file_path: str) -> ExtractedText:
    """Parse a DOCX file and return extracted text asynchronously.

    Parameters
    ----------
    file_path: str
        Path to the DOCX file.

    Returns
    -------
    ExtractedText
        Structured text extracted from the DOCX document.
    """
    # Run the synchronous operation in a thread pool
    loop = asyncio.get_event_loop()

    def _parse_sync():
        from docx import Document  # Imported lazily
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = Document(file_path)
        except PackageNotFoundError as e:
            raise FileNotFoundError(f"DOCX file not found: {file_path}") from e

        # Extract text from paragraphs
        text_parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        text = "\n".join(text_parts)
        return text

    text = await loop.run_in_executor(None, _parse_sync)

    pages = [PageText(page_number=1, text=text, ocr=False)]
    return ExtractedText(text=text, file_type="docx", pages=pages)
